import re
import pandas as pd
import jinja2
from jinja2 import Environment, BaseLoader, StrictUndefined
from google.genai import types
from docxtpl import InlineImage
from docx.shared import Mm, Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
from datetime import date

def detect_header(filepath, sheet_name=0, required_columns=None):
    """
    Detecta a linha de cabeçalho em um arquivo Excel procurando pelas colunas obrigatórias.
    Retorna o índice da linha (0-based) ou levanta ValueError se não encontrar.
    """
    if not required_columns:
        return 0 # Default to 0 if no columns specified

    # Read a sample of rows to search for header
    try:
        # Ler as primeiras 20 linhas sem cabeçalho para inspecionar
        df_sample = pd.read_excel(filepath, sheet_name=sheet_name, header=None, nrows=20)
    except Exception as e:
        filename = getattr(filepath, 'name', str(filepath))
        raise ValueError(f"Erro ao ler amostra do arquivo '{filename}' aba '{sheet_name}': {e}")

    required_set = set(c.lower().strip() for c in required_columns)

    for i, row in df_sample.iterrows():
        # Convert row values to strings, lowercased and stripped, filtering out NaNs
        row_values = set(str(val).lower().strip() for val in row.values if pd.notna(val))

        # Check if all required columns are present in this row
        if required_set.issubset(row_values):
            return i

    filename = getattr(filepath, 'name', str(filepath))
    raise ValueError(f"Não foi possível detectar o cabeçalho na aba '{sheet_name}' do arquivo '{filename}'. Colunas esperadas: {', '.join(required_columns)}")

def validar_schema(df, required_columns):
    """
    Valida se as colunas obrigatórias estão presentes no DataFrame.
    Lança ValueError se faltar alguma.
    """
    if not required_columns:
        return

    df_cols = set(c.lower().strip() for c in df.columns)
    req_cols = set(c.lower().strip() for c in required_columns)

    missing = req_cols - df_cols
    if missing:
        raise ValueError(f"Colunas obrigatórias ausentes: {', '.join(missing)}")

def carregar_dados(filepath, sheet_name=0, skiprows=None, required_columns=None):
    """
    Lê um arquivo Excel e retorna um DataFrame.

    Args:
        filepath: Caminho ou objeto do arquivo.
        sheet_name: Nome ou índice da aba.
        skiprows: (Opcional) Número de linhas para pular. Se None e required_columns for fornecido, tenta detectar.
        required_columns: (Opcional) Lista de nomes de colunas obrigatórias para detecção e validação.
    """
    try:
        # Se skiprows não for informado mas tivermos colunas obrigatórias, tenta detectar
        if skiprows is None and required_columns:
            skiprows = detect_header(filepath, sheet_name, required_columns)
        elif skiprows is None:
            skiprows = 0 # Default legacy behavior

        # Carrega o dataframe final
        df = pd.read_excel(filepath, sheet_name=sheet_name, skiprows=skiprows)

        # Normaliza nomes das colunas para strip (remove espaços extras)
        df.columns = [str(c).strip() for c in df.columns]

        # Aplica map apenas nas células de string
        df = df.map(lambda x: x.strip() if isinstance(x, str) else x)

        # Valida schema se solicitado
        if required_columns:
            # A validação aqui é case-insensitive por segurança,
            # mas idealmente os nomes devem bater exato ou normalizarmos tudo.
            # Para garantir robustez, vamos validar contra os nomes normalizados (strip).
            # A função validar_schema já faz lower().strip() para comparar.
            validar_schema(df, required_columns)

        return df

    except Exception as e:
        filename = getattr(filepath, 'name', str(filepath))
        raise ValueError(f"Erro ao carregar a planilha '{sheet_name}' do arquivo '{filename}': {e}")


def aplicar_variaveis_temporarias(fontes, variaveis):
    """Materializa colunas derivadas nas fontes apenas durante o processamento."""
    if variaveis is None or variaveis.empty:
        return fontes

    required = {'id', 'id_fonte_informacao', 'nome', 'expressao'}
    missing = required - set(variaveis.columns)
    if missing:
        raise ValueError(
            "Colunas obrigatórias ausentes em Variáveis Temporárias: "
            + ", ".join(sorted(missing))
        )

    nomes_criados = set()
    for _, row in variaveis.iterrows():
        variable_id = str(row['id']).strip()
        source_id = str(row['id_fonte_informacao']).strip()
        name = str(row['nome']).strip()
        expression = str(row['expressao']).strip()
        source = fontes.get(source_id)

        if source is None:
            raise ValueError(
                f"Variável temporária {variable_id} refere-se à fonte inexistente {source_id!r}."
            )
        if source.info is None:
            raise ValueError(
                f"A fonte {source_id!r} não foi carregada para calcular a variável {variable_id}."
            )
        if not name or name.lower() == 'nan':
            raise ValueError(f"Variável temporária {variable_id} sem nome.")
        if name in nomes_criados or name in source.info.columns:
            raise ValueError(f"Nome de variável temporária duplicado ou já existente: {name!r}.")

        prepared = expression
        for column in sorted(source.info.columns, key=lambda value: len(str(value)), reverse=True):
            column = str(column)
            if not re.fullmatch(r"[A-Za-z_]\w*", column):
                prepared = prepared.replace(column, f"`{column}`")

        try:
            source.info[name] = source.info.eval(prepared, engine='python')
        except Exception as exc:
            raise ValueError(
                f"Não foi possível calcular a variável temporária {variable_id} ({name}): {exc}"
            ) from exc
        nomes_criados.add(name)

    return fontes

def get_variaveis_template(template_md_content):
    """Coleta as variáveis presentes em um template Jinja2."""
    if not template_md_content:
        return set()
    env = Environment(loader=BaseLoader())
    ast = env.parse(template_md_content)
    return jinja2.meta.find_undeclared_variables(ast)

class StreamlitLogHandler(logging.Handler):
    """Handler de logging customizado para exibir logs do pypandoc no Streamlit."""
    def __init__(self, container):
        super().__init__()
        self.container = container
        self.records = []

    def emit(self, record):
        self.records.append(record)
        msg = self.format(record)
        self.container.warning(msg)


def parse_expression(expression):
    tokens = []
    current_token = ''

    for char in expression:
        if char in {'|', '&', '~', '(', ')'}:
            if current_token:
                tokens.append(current_token.strip())
                current_token = ''
            tokens.append(char)
        else:
            current_token += char

    if current_token:
        tokens.append(current_token.strip())

    # Remove empty tokens that might result from spaces like "A | B" -> "A", "", "|", "", "B"
    return [t for t in tokens if t]

def infix_to_rpn(tokens):
    precedence = {'~': 3, '&': 2, '|': 1}
    output = []
    stack = []

    for token in tokens:
        if token == '(':
            # Abre parêntese
            stack.append(token)
        elif token == ')':
            # Fecha parêntese
            while stack and stack[-1] != '(':
                output.append(stack.pop())
            stack.pop()  # Remove o '('
        elif token in precedence:
            # Operador
            while stack and stack[-1] in precedence and precedence[stack[-1]] >= precedence[token]:
                output.append(stack.pop())
            stack.append(token)
        else:
            output.append(token)

    while stack:
        output.append(stack.pop())

    return list(filter(lambda x: x != '', output))

def safe_compare(value, condition):
    """
    Compara um valor com uma condição de forma segura (sem eval).
    Suporta: >, <, >=, <=, ==, != e igualdade implícita.
    """
    condition = str(condition).strip()
    value_str = str(value).strip()

    # Operadores suportados
    operators = {
        '>=': lambda x, y: x >= y,
        '<=': lambda x, y: x <= y,
        '!=': lambda x, y: x != y,
        '==': lambda x, y: x == y,
        '>': lambda x, y: x > y,
        '<': lambda x, y: x < y
    }

    # Tenta encontrar um operador no início da condição
    # Ordenamos chaves pelo tamanho decrescente para pegar >= antes de >
    for op_symbol in sorted(operators.keys(), key=len, reverse=True):
        if condition.startswith(op_symbol):
            target_str = condition[len(op_symbol):].strip()
            # Remove quotes if user added them for a number (e.g. > '5')
            target_clean = target_str.strip("'\"")

            # Tenta comparação numérica
            try:
                val_num = float(value)
                target_num = float(target_clean) # Convert cleaned target
                return operators[op_symbol](val_num, target_num)
            except ValueError:
                # Comparação de strings (remove aspas se houver)
                # target_clean já está limpo
                return operators[op_symbol](value_str, target_clean)

    # Se não houver operador, assume igualdade (==)
    # Remove aspas se o usuário tiver colocado (ex: 'Sim')
    target_clean = condition.strip("'\"")
    # Tenta comparar como número se ambos forem numéricos e a condição for apenas um número
    try:
        val_num = float(value)
        target_num = float(target_clean)
        return val_num == target_num
    except ValueError:
        pass

    return value_str == target_clean

def avalia_logica(expressao, contexto):
    """
    Avalia uma expressão lógica (ex: 'AV01 & AV02') dado um contexto
    (dicionário de resultados { 'AV01': True, ... }).
    """
    if not expressao:
        return False

    tokens = parse_expression(str(expressao))
    rpn = infix_to_rpn(tokens)

    pilha = []

    for token in rpn:
        if token == '|':
            if len(pilha) < 2: raise ValueError("Expressão mal formada (operador |)")
            y = pilha.pop()
            x = pilha.pop()
            pilha.append(x or y)
        elif token == '&':
            if len(pilha) < 2: raise ValueError("Expressão mal formada (operador &)")
            y = pilha.pop()
            x = pilha.pop()
            pilha.append(x and y)
        elif token == '~':
            if len(pilha) < 1: raise ValueError("Expressão mal formada (operador ~)")
            x = pilha.pop()
            pilha.append(not x)
        else:
            # Token é uma chave no contexto (ex: AV01)
            val = contexto.get(token, False) # Default False se não achar
            pilha.append(val)

    if len(pilha) == 1:
        return pilha[0]
    else:
        # Se sobrou mais de um item, expressão pode estar incompleta, mas retornamos o topo
        return pilha[0]

def avalia_expressao(expressao_achado, situacao_encontrada, debug=False):
    expressao_achado = str(expressao_achado)
    # situacao_encontrada pode ser int/float/str, mantemos o tipo original para safe_compare tentar converter

    # Alternativas exportadas pelo LimeSurvey podem conter marcadores como
    # "f)" e outros parênteses que fazem parte do valor literal. Nesse caso,
    # somente operadores separados por espaços e os delimitadores externos
    # são estruturais; a pontuação interna deve permanecer na comparação.
    if re.search(r"(?:^|[~(])\s*[a-z]\)\s", expressao_achado, flags=re.IGNORECASE):
        def avaliar_alternativa(expressao):
            expressao = expressao.strip()
            if expressao.startswith('~(') and expressao.endswith(')'):
                return not avaliar_alternativa(expressao[2:-1])
            if (
                expressao.startswith('(')
                and expressao.endswith(')')
                and (' | ' in expressao or ' & ' in expressao)
            ):
                expressao = expressao[1:-1].strip()
            if ' | ' in expressao:
                return any(avaliar_alternativa(item) for item in expressao.split(' | '))
            if ' & ' in expressao:
                return all(avaliar_alternativa(item) for item in expressao.split(' & '))
            if expressao.startswith('~'):
                return not avaliar_alternativa(expressao[1:])
            return safe_compare(situacao_encontrada, expressao)

        resultado = avaliar_alternativa(expressao_achado)
        if debug:
            print(f"Avalia alternativa literal: '{expressao_achado}' vs '{situacao_encontrada}' -> {resultado}")
        return resultado

    parsed_tokens = parse_expression(expressao_achado)
    tokens = infix_to_rpn(parsed_tokens)

    if debug:
        print(f"Avalia Expressão: '{expressao_achado}' vs '{situacao_encontrada}'")
        print(f"Tokens RPN: {tokens}")

    pilha = []

    for token in tokens:
        if token == '|':
            if len(pilha) < 2: raise ValueError("Expressão mal formada (operador |)")
            y = pilha.pop()
            x = pilha.pop()
            pilha.append(x or y)
        elif token == '&':
            if len(pilha) < 2: raise ValueError("Expressão mal formada (operador &)")
            y = pilha.pop()
            x = pilha.pop()
            pilha.append(x and y)
        elif token == '~':
            if len(pilha) < 1: raise ValueError("Expressão mal formada (operador ~)")
            x = pilha.pop()
            pilha.append(not x)
        else:
            # O token é uma condição simples (ex: "> 10" ou "Sim")
            # Removemos o '.' no final que às vezes vem da planilha
            condition = re.sub(r'\.$', '', token)

            # Usamos safe_compare para validar
            resultado = safe_compare(situacao_encontrada, condition)

            if debug:
                print(f"    Comparando: Val='{situacao_encontrada}' Cond='{condition}' -> {resultado}")

            pilha.append(resultado)

    if len(pilha) >= 1:
        if debug:
            print('        Resultado Final:', pilha[0])
        return pilha[0]
    else:
        raise ValueError("Expressão lógica inválida ou vazia")

def processa_imagens_contexto(contexto, context_files_path_map, template_type, base_docx=None):
    """
    Substitui nomes de arquivos de imagem no contexto pelos caminhos ou objetos de imagem apropriados.

    Returns:
        tuple: (contexto_atualizado, lista_de_avisos)
    """
    image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp')
    warnings = []

    # Itera sobre uma cópia dos itens para permitir a modificação do dicionário
    for key, value in list(contexto.items()):
        if isinstance(value, str) and value.lower().endswith(image_extensions):
            if value in context_files_path_map:
                image_path = context_files_path_map[value]
                if template_type == 'docx':
                    if base_docx is None:
                        raise ValueError("O objeto base_docx é necessário para processar imagens em templates .docx")
                    # Para docx, substitui pelo objeto InlineImage
                    contexto[key] = InlineImage(base_docx, image_path, width=Mm(160))
                elif template_type == 'md':
                    # Para markdown, substitui pelo caminho do arquivo
                    contexto[key] = image_path
            else:
                warnings.append(f"Arquivo de imagem '{value}' para a variável '{key}' não encontrado. A imagem não será inserida.")
                contexto[key] = f"[Imagem '{value}' não encontrada]"

    return contexto, warnings

def cross_ref_figuras(template_str: str) -> str:
    """
    Processa um template de texto para numerar automaticamente as referências
    de figuras e modificar as legendas das imagens, preservando os
    atributos de formatação do Pandoc (ex: {width=10cm}).

    A função opera em duas passadas:

    1. Mapeamento:
       Encontra todas as declarações ({#fig:ID#}) e referências ([@fig:ID])
       na ordem em que aparecem para criar um mapa de numeração
       (ex: {'fig_01': 1, 'fig_02': 2}).

    2. Substituição:
       - Substitui referências de texto (ex: [@fig:fig_01] -> "Figura 1").
       - Encontra as linhas de imagem (ex: ![Texto]({{path}}){attrs}{#fig:ID#})
         e as substitui por "![Figura 1 - Texto]({{path}}){attrs}".
    """

    # --- Passa 1: Mapeamento ---

    figura_map = {}
    contador = 1

    # Regex combinada (NÃO MUDA)
    regex_combinado = r"(?:\{#fig:([^#]+)#\}|\[@fig:([^\]]+)\])"

    for match in re.finditer(regex_combinado, template_str):
        id_declaracao = match.group(1)
        id_referencia = match.group(2)
        fig_id = id_declaracao if id_declaracao else id_referencia

        if fig_id and fig_id not in figura_map:
            figura_map[fig_id] = contador
            contador += 1

    if not figura_map:
        return template_str

    # --- Passa 2: Substituições ---

    texto_processado = template_str

    # 1. Substituir referências de TEXTO (NÃO MUDA)
    regex_ref_texto = r"\[@fig:([^\]]+)\]"

    def substituir_ref_texto(match):
        fig_id = match.group(1)
        if fig_id in figura_map:
            return f"Figura {figura_map[fig_id]}"
        return match.group(0)

    texto_processado = re.sub(regex_ref_texto, substituir_ref_texto, texto_processado)

    # 2. Modificar linhas de IMAGEM e remover tags de declaração (MODIFICADO)

    # Regex ATUALIZADA:
    # Grupo 1: ![ (alt text) ]
    # Grupo 2: (path)
    # Grupo 3: (bloco de atributos opcional, ex: {width=10cm})
    # Grupo 4: {#fig: (ID) #}
    regex_imagem_decl = r"!\[([^\]]*)\](\([^)]*\))\s*(\{[^}]*\})?\s*\{#fig:([^#]+)#\}"

    def modificar_legenda_imagem(match):
        alt_text = match.group(1)
        path = match.group(2)
        attributes = match.group(3)  # O bloco {width=10cm}
        fig_id = match.group(4)

        if fig_id in figura_map:
            numero = figura_map[fig_id]

            # Se o grupo de atributos não for encontrado (None),
            # o transformamos em uma string vazia.
            attr_str = attributes if attributes else ""

            # Reconstrói a string: ![Figura X - Texto](path){atributos}
            return f"![Figura {numero} - {alt_text}]{path}{attr_str}"

        return match.group(0) # Failsafe

    texto_processado = re.sub(regex_imagem_decl, modificar_legenda_imagem, texto_processado)

    return texto_processado

def processar_quebras_pagina(template_str: str) -> str:
    """
    Substitui o comando LaTeX \newpage por um bloco Raw OpenXML
    que o Pandoc entende e converte corretamente para quebra de página no Word.
    """
    # Bloco nativo do OpenXML para quebra de página
    pagebreak_openxml = "\n```{=openxml}\n<w:p><w:r><w:br w:type=\"page\"/></w:r></w:p>\n```\n"

    # Regex para encontrar \newpage (aceitando espaços extras ou chaves vazias opcionais)
    regex_newpage = r"\\newpage(?:\{\})?"

    return re.sub(regex_newpage, pagebreak_openxml, template_str)

def cross_ref_tabelas(template_str: str) -> str:
    """
    Processa um template de texto para numerar automaticamente as referências
    de tabelas e modificar as legendas das tabelas, seguindo o padrão de
    cross_ref_figuras.
    """

    # --- Passa 1: Mapeamento ---

    tabela_map = {}
    contador = 1

    # Regex combinada para encontrar declarações e referências
    regex_combinado = r"(?:\{#tbl:([^#]+)#\}|\[@tbl:([^\]]+)\])"

    for match in re.finditer(regex_combinado, template_str):
        id_declaracao = match.group(1)
        id_referencia = match.group(2)
        tbl_id = id_declaracao if id_declaracao else id_referencia

        if tbl_id and tbl_id not in tabela_map:
            tabela_map[tbl_id] = contador
            contador += 1

    if not tabela_map:
        return template_str

    # --- Passa 2: Substituições ---

    texto_processado = template_str

    # 1. Substituir referências de TEXTO
    regex_ref_texto = r"\[@tbl:([^\]]+)\]"

    def substituir_ref_texto(match):
        tbl_id = match.group(1)
        if tbl_id in tabela_map:
            return f"Tabela {tabela_map[tbl_id]}"
        return match.group(0)

    texto_processado = re.sub(regex_ref_texto, substituir_ref_texto, texto_processado)

    # 2. Modificar legendas de Tabela
    # Padrão esperado: ": Legenda da Tabela {#tbl:ID#}" OU "Table: Legenda..."
    # (?m) habilita multiline para ^ coincidir com início da linha
    regex_legenda = r"(?m)^(:|Table:)\s*(.*?)\s*\{#tbl:([^#]+)#\}"

    def modificar_legenda_tabela(match):
        prefix = match.group(1) # ":" ou "Table:"
        caption = match.group(2)
        tbl_id = match.group(3)

        if tbl_id in tabela_map:
            numero = tabela_map[tbl_id]
            # Reconstrói mantendo o formato que o Pandoc espera para criar a legenda corretamente
            return f"{prefix} Tabela {numero} - {caption} {{#tbl:{tbl_id}#}}"

        return match.group(0)

    texto_processado = re.sub(regex_legenda, modificar_legenda_tabela, texto_processado)

    return texto_processado


def substituir_underline_preview(text: str) -> str:
    """
    Substitui texto envolvido por __ pela sintaxe de underline HTML:
    __texto__ -> <u>texto</u>
    """
    return re.sub(r"__(.+?)__", r"<u>\1</u>", text)


def substituir_underline_pandoc(text: str) -> str:
    """
    Substitui texto envolvido por __ pela sintaxe de underline do Pandoc:
    __texto__ -> [texto]{.underline}
    """
    return re.sub(r"__(.+?)__", r"[\1]{.underline}", text)


def avalia_gemini(client, prompt_text: str, modelo, temperature, response_format_choice, file_objects = []):
    """
    Chama a API do Gemini com a configuração apropriada.
    Retorna a resposta do modelo e uma mensagem de erro (se houver).
    """
    try:
        contents = [prompt_text] + file_objects
        # st.info(f'Avaliando com o modelo {modelo}...') # Comentado para evitar chamadas Streamlit em utils

        generation_config = types.GenerateContentConfig(
            max_output_tokens=65536, # Usando o valor sugerido de 65536
            temperature=temperature,
        )

        if response_format_choice == 'Estruturada':
            generation_config.response_mime_type = 'application/json'

        # Cria o conteúdo para a API
        response = client.models.generate_content(
            model=modelo,
            contents=contents,
            config=generation_config
        )

        return response, None

    except Exception as e:
        return None, f"Erro ao chamar a API Gemini: {e}"

def data_hoje_abnt():
    hoje = date.today()

    # Lista de meses ABNT (Jan. Fev. Mar. Abr. Maio Jun. Jul. Ago. Set. Out. Nov. Dez.)
    # Note que 'maio' não tem ponto e é minúsculo na citação direta,
    # mas aqui usaremos minúsculo padrão.
    meses_abnt = {
        1: 'jan.', 2: 'fev.', 3: 'mar.', 4: 'abr.', 5: 'maio', 6: 'jun.',
        7: 'jul.', 8: 'ago.', 9: 'set.', 10: 'out.', 11: 'nov.', 12: 'dez.'
    }

    dia = hoje.day
    mes = meses_abnt[hoje.month]
    ano = hoje.year

    return f"{dia} {mes} {ano}"

def data_hoje():
    return date.today().strftime("%d/%m/%Y")

def aplicar_estilo_tabelas(docx_path, font_name='Calibri', header_size=10, body_size=9):
    """
    Aplica estilos específicos às tabelas de um arquivo DOCX:
    - Cabeçalho: Fonte parametrizada (padrão Calibri), tamanho parametrizado (padrão 10)
    - Corpo: Fonte parametrizada (padrão Calibri), tamanho parametrizado (padrão 9)
    - Alinhamento: Justificado para todo o conteúdo
    """
    doc = Document(docx_path)
    for table in doc.tables:
        # Aplicar "ajustar-se automaticamente ao conteúdo"
        table.autofit = True # or WD_TABLE_AUTOFORMAT.AUTOFIT_CONTENTS if needed but True is often sufficient

        for i, row in enumerate(table.rows):
            is_header = (i == 0)
            font_size = Pt(header_size) if is_header else Pt(body_size)

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = font_size

    doc.save(docx_path)

def apply_custom_style():
    """
    Aplica estilos CSS customizados para tornar a interface mais moderna e profissional.
    Deve ser chamado no início de cada página Streamlit.
    """
    import streamlit as st

    st.markdown("""
        <style>
        /* Import Google Fonts */
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');

        /*
        html, body, [class*="css"] {
            font-family: 'Inter', sans-serif;
        }*/

        /* Títulos */
        h1 {
            color: #004e92;
            font-weight: 700 !important;
            letter-spacing: -0.5px;
        }
        h2, h3 {
            color: #2c3e50;
            font-weight: 600 !important;
        }

        /* Sidebar Styling */
        /*section[data-testid="stSidebar"] {
            background-color: #f8f9fa;
            border-right: 1px solid #e9ecef;
        }*/

        /* Cards customizados (Simulação com markdown) */
        .card {
            background-color: white;
            padding: 1.5rem;
            border-radius: 10px;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.05);
            margin-bottom: 1rem;
            border: 1px solid #e9ecef;
            transition: transform 0.2s ease;
        }
        .card:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 12px rgba(0, 0, 0, 0.1);
        }
        .card h4 {
            color: #0068C9;
            margin-top: 0;
            font-weight: 600;
        }
        .card p {
            color: #6c757d;
            font-size: 0.95rem;
            margin-bottom: 0;
        }

        /* Botões */
        div.stButton > button {
            background-color: #0068C9;
            color: white;
            border-radius: 8px;
            padding: 0.5rem 1rem;
            font-weight: 600;
            border: none;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            transition: all 0.2s;
        }
        div.stButton > button:hover {
            background-color: #0056b3;
            box-shadow: 0 4px 8px rgba(0,0,0,0.15);
        }
        div.stButton > button:active {
            background-color: #004494;
        }

        /* File Uploader */
        section[data-testid="stFileUploader"] {
            border-radius: 10px;
            padding: 1rem;
            border: 1px dashed #ced4da;
            background-color: #fdfdfe;
        }

        /* Expander */
        .streamlit-expanderHeader {
            font-weight: 600;
            color: #495057;
            background-color: white;
            border-radius: 8px;
        }

        /* Mensagens de Sucesso/Info/Erro */
        .stAlert {
            border-radius: 8px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }

        </style>
    """, unsafe_allow_html=True)

def create_card(title, content, icon=None):
    """Retorna o HTML para um card estilizado."""
    icon_html = f'<div style="font-size: 1.5rem; margin-bottom: 0.5rem;">{icon}</div>' if icon else ''
    return f"""
    <div class="card">
        {icon_html}
        <h4>{title}</h4>
        <p>{content}</p>
    </div>
    """
