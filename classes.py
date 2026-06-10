import os
import re
import copy
import pandas as pd
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from utils import avalia_expressao, avalia_logica

def safe_serialize(obj):
    """Helper to serialize numpy/pandas types to native Python types."""
    if isinstance(obj, (np.integer, np.int64)):
        return int(obj)
    elif isinstance(obj, (np.floating, np.float64)):
        return float(obj) if not pd.isna(obj) else None
    elif isinstance(obj, (np.bool_, bool)):
        return bool(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif pd.isna(obj):
        return None
    return obj

class FonteInformacao:
    contador = 1  # Contador de instâncias para automatizar o identificador

    def __init__(self, descricao, filepath, chave_jurisdicionado=None, id=None):
        if id is None:
            id = f"FI{FonteInformacao.contador:02d}"
            FonteInformacao.contador += 1  # Incrementa o contador para o próximo identificador

        self.id = id
        self.descricao = descricao  # Descrição da fonte de informação
        self.filepath = filepath
        self.chave_jurisdicionado = chave_jurisdicionado
        self.info = None

    def __repr__(self):
        return f"FonteInformacao(id='{self.id}', descricao='{self.descricao}', filepath='{self.filepath}', chave_jurisdicionado='{self.chave_jurisdicionado}')"

    def read(self):
        """Lê o conteúdo da fonte de informação, assumindo que seja uma planilha Excel."""
        try:
            # Tenta ler uma amostra do arquivo para verificar se é uma planilha
            pd.read_excel(self.filepath, nrows=1)
            self.info = pd.read_excel(self.filepath)
            if self.chave_jurisdicionado:
                # Trata chaves duplicadas e remove espaços em branco extras
                if self.chave_jurisdicionado in self.info.columns:
                    self.info[self.chave_jurisdicionado] = self.info[self.chave_jurisdicionado].astype(str).str.strip()
                    # Se houver coluna de data de envio, ordena para priorizar o enviado mais recentemente
                    if 'submitdate' in self.info.columns:
                        self.info = self.info.sort_values(by='submitdate', na_position='last')
                    # Remove duplicatas baseadas na chave do jurisdicionado
                    self.info = self.info.drop_duplicates(subset=[self.chave_jurisdicionado], keep='first')
                self.info = self.info.set_index(self.chave_jurisdicionado)
        except Exception as e:
            if isinstance(self.filepath, str):
                msg = f"Arquivo '{self.filepath}' não pôde ser lido como planilha Excel: {e}"
            else:
                msg = f"O arquivo carregado não pôde ser lido como planilha Excel: {e}"
            raise IOError(msg)

    def to_dict(self):
        # Handle Streamlit UploadedFile objects which are not JSON serializable
        filepath_val = self.filepath
        if hasattr(filepath_val, 'name'):
            filepath_val = filepath_val.name

        return {
            'id': self.id,
            'descricao': self.descricao,
            'filepath': filepath_val,
            'chave_jurisdicionado': self.chave_jurisdicionado
        }

    @classmethod
    def from_dict(cls, data):
        return cls(
            id=data.get('id'),
            descricao=data.get('descricao'),
            filepath=data.get('filepath'),
            chave_jurisdicionado=data.get('chave_jurisdicionado')
        )

class Achado:
    def __init__(self, numero, nome, situacoes_encontradas=None, evidencias=None, encaminhamentos=None):
        self.numero = numero
        self.nome = nome
        self.situacoes_encontradas = situacoes_encontradas if situacoes_encontradas is not None else []
        self.encaminhamentos = encaminhamentos if encaminhamentos is not None else []
        self.evidencias = evidencias if evidencias is not None else []

    def __repr__(self):
        return  f"Achado(numero='{self.numero}', nome='{self.nome}')"

    def to_dict(self):
        return {
            'numero': self.numero,
            'nome': self.nome,
            'situacoes_encontradas': [safe_serialize(s) for s in self.situacoes_encontradas],
            'evidencias': [safe_serialize(e) for e in self.evidencias],
            'encaminhamentos': self.encaminhamentos
        }

    @classmethod
    def from_dict(cls, data):
        return cls(
            numero=data.get('numero'),
            nome=data.get('nome'),
            situacoes_encontradas=data.get('situacoes_encontradas'),
            evidencias=data.get('evidencias'),
            encaminhamentos=data.get('encaminhamentos')
        )

class AcaoVerificacao:
    contador = 1  # Contador de instâncias para automatizar o identificador

    def __init__(self, fonte_informacao, informacao_requerida, descricao_evidencia, situacao_inconforme,
                 tipo_encaminhamento, encaminhamento, pre_encaminhamento, criterio, descricao_situacao_inconforme,
                 acao_exclusiva_auditados=None, auditado_inexistente_e_achado=None,
                 descricao_auditado_inexistente=None, situacao_encontrada_nan_e_achado=None, id=None):

        if id is None:
            id = f"AV{AcaoVerificacao.contador:02d}"
            AcaoVerificacao.contador += 1  # Incrementa o contador para o próximo identificador

        self.id = id
        self.fonte_informacao = fonte_informacao  # Objeto FonteInformacao associado
        self.informacao_requerida = informacao_requerida  # Campo/coluna a ser verificada

        self.criterio = criterio  # Condição a ser avaliada (ex: "> 0")
        self.descricao_evidencia = descricao_evidencia  # Descrição da evidência em caso de inconformidade

        self.acao_exclusiva_auditados = acao_exclusiva_auditados  # Siglas dos auditados (CSV)
        self.auditado_inexistente_e_achado = auditado_inexistente_e_achado  # Se auditado não existir, reporta achado?
        self.descricao_auditado_inexistente = descricao_auditado_inexistente  # Descrever o pq não existe

        self.situacao_inconforme = situacao_inconforme  # Condição que indica inconformidade (ex: "Não adota")
        self.descricao_situacao_inconforme = descricao_situacao_inconforme
        self.situacao_encontrada_nan_e_achado = False if pd.isna(situacao_encontrada_nan_e_achado) else True
        self.situacao_encontrada = None  # Valor encontrado durante a verificação

        self.tipo_encaminhamento = tipo_encaminhamento  # Ex: "Recomendação", "Determinação"
        self.pre_encaminhamento = pre_encaminhamento  # Ação prévia ao encaminhamento
        self.encaminhamento = encaminhamento

        self.resultado = None    # Inicialmente falso, até ser verificado


    def __repr__(self):
        return (
                f"AcaoVerificacao(id='{self.id}', fonte_informacao='{self.fonte_informacao}', \n"
                f"informacao_requerida='{self.informacao_requerida}', criterio='{self.criterio}', descricao_evidencia='{self.descricao_evidencia}')\n"
                f"situacao_inconforme='{self.situacao_inconforme}', tipo_encaminhamento='{self.tipo_encaminhamento}')\n"
                f"descricao_situacao_inconforme='{self.descricao_situacao_inconforme}', acao_exclusiva_auditados='{self.acao_exclusiva_auditados}')\n"
                f"pre_encaminhamento='{self.pre_encaminhamento}', encaminhamento='{self.encaminhamento}')\n"
                f"situacao_encontrada_nan_e_achado='{self.situacao_encontrada_nan_e_achado}', auditado_inexistente_e_achado='{self.auditado_inexistente_e_achado}')\n"
                f"descricao_auditado_inexistente='{self.descricao_auditado_inexistente}', situacao_encontrada='{self.situacao_encontrada}')\n"
                f"resultado='{self.resultado}'")

    def executar(self, auditado, debug=False):
        # Verifica se a ação é exclusiva para um determinado grupo de auditados.
        # Se for, e o auditado atual não estiver nesse grupo, a ação não é executada.
        # Isso permite que certas verificações sejam feitas apenas em alguns órgãos.
        if not pd.isna(self.acao_exclusiva_auditados) and (auditado not in self.acao_exclusiva_auditados):
            return self

        if debug:
            print(f'\tExecutando ação {self.id}')
            print(f'\tBuscar em "{self.fonte_informacao.descricao}" no campo "{self.informacao_requerida}" se ocorre "{self.situacao_inconforme}".')

        # Verifica se a busca será feita em mais de um campo específico na fonte de informação
        for info_requerida in self.informacao_requerida.split('|'):
            if info_requerida not in self.fonte_informacao.info.columns:
                print(f'ERROR: Não foi possível encontrar o campo "{self.informacao_requerida}" na fonte de informação.')
                print(f'ERROR: Ajuste o campo ou a fonte.')
                # return self
                raise ValueError(f'Na Ação de Verificação "{self.id}", não foi possível encontrar a coluna "{info_requerida}" na fonte de informação "{self.fonte_informacao.descricao}". '
                                 f'Verifique se o nome da coluna está correto no "mapa-verificacao-achados.xlsx".')

        # Realiza a busca e a verificação para cada campo especificado
        if auditado in self.fonte_informacao.info.index:
            resultado_acoes = []
            for info_requerida in self.informacao_requerida.split('|'):
                val = self.fonte_informacao.info.loc[auditado, info_requerida]
                # Defesa contra índices duplicados que retornam uma Series ao invés de um valor escalar
                if isinstance(val, pd.Series):
                    val = val.iloc[0]
                self.situacao_encontrada = val

                if self.situacao_encontrada_nan_e_achado and pd.isna(self.situacao_encontrada):
                    resultado_acoes.append(True)
                else:
                    resultado_acoes.append(avalia_expressao(self.situacao_inconforme, self.situacao_encontrada, debug=debug))
                    self.descricao_evidencia = self.descricao_evidencia.replace('@', str(self.situacao_encontrada))

            self.resultado = all(resultado_acoes)

        else:
            self.resultado = True if self.auditado_inexistente_e_achado else False
            self.descricao_evidencia = self.descricao_auditado_inexistente

        if debug:
            print(f'\tSituação Encontrada: {self.situacao_encontrada}')
            print(f'\tResultado da verificação: {self.resultado}')
            print(f'')

        return self

    def to_dict(self):
        return {
            'id': self.id,
            'fonte_informacao': self.fonte_informacao.to_dict(),
            'informacao_requerida': self.informacao_requerida,
            'criterio': self.criterio,
            'descricao_evidencia': self.descricao_evidencia,
            'situacao_inconforme': self.situacao_inconforme,
            'tipo_encaminhamento': self.tipo_encaminhamento,
            'encaminhamento': self.encaminhamento,
            'pre_encaminhamento': self.pre_encaminhamento,
            'descricao_situacao_inconforme': self.descricao_situacao_inconforme,
            'acao_exclusiva_auditados': safe_serialize(self.acao_exclusiva_auditados),
            'auditado_inexistente_e_achado': safe_serialize(self.auditado_inexistente_e_achado),
            'descricao_auditado_inexistente': self.descricao_auditado_inexistente,
            'situacao_encontrada_nan_e_achado': safe_serialize(self.situacao_encontrada_nan_e_achado),
            'situacao_encontrada': safe_serialize(self.situacao_encontrada),
            'resultado': bool(self.resultado) if self.resultado is not None else None
        }

    @classmethod
    def from_dict(cls, data):
        obj = cls(
            id=data.get('id'),
            fonte_informacao=FonteInformacao.from_dict(data.get('fonte_informacao')),
            informacao_requerida=data.get('informacao_requerida'),
            criterio=data.get('criterio'),
            descricao_evidencia=data.get('descricao_evidencia'),
            situacao_inconforme=data.get('situacao_inconforme'),
            tipo_encaminhamento=data.get('tipo_encaminhamento'),
            encaminhamento=data.get('encaminhamento'),
            pre_encaminhamento=data.get('pre_encaminhamento'),
            descricao_situacao_inconforme=data.get('descricao_situacao_inconforme'),
            acao_exclusiva_auditados=data.get('acao_exclusiva_auditados'),
            auditado_inexistente_e_achado=data.get('auditado_inexistente_e_achado'),
            descricao_auditado_inexistente=data.get('descricao_auditado_inexistente'),
            situacao_encontrada_nan_e_achado=data.get('situacao_encontrada_nan_e_achado')
        )
        obj.situacao_encontrada = data.get('situacao_encontrada')
        obj.resultado = data.get('resultado')
        return obj

class ProcedimentoAuditoria:
    contador = 1  # Contador de instâncias para automatizar o identificador

    def __init__(self, descricao, logica_achado, numero_achado, nome_achado, id=None):
        if id is None:
            id = f"PA{ProcedimentoAuditoria.contador:02d}"
            ProcedimentoAuditoria.contador += 1  # Incrementa o contador para o próximo identificador

        self.id = id
        self.descricao = descricao  # Descrição do procedimento
        self.logica_achado = logica_achado  # Expressão lógica para determinar o achado (ex: "AV01 | AV02")
        self.executado = False

        self.numero_achado = numero_achado
        self.nome_achado = nome_achado

        # Ações de verificação que compõem o procedimento
        self.acoes_verificacao = []
        self.achado = None
        self.achado_ocorreu = None

    def __repr__(self):
        return  (f"ProcedimentoAuditoria(id='{self.id}', \n" +
                f"descricao='{self.descricao}'\n" +
                f"executado='{self.executado}'\n" +
                f"logica_achado='{self.logica_achado}'\n" +
                f"achado_ocorreu='{self.achado_ocorreu}'\n" +
                f"achado='{self.achado}'\n" +
                f"acoes_verificacao ('{len(self.acoes_verificacao)}')\n") #+ "\n".join([f"{acao}" for acao in self.acoes_verificacao])


    def adicionar_acao(self, acao):
        """Adiciona uma ação de verificação ao procedimento."""
        self.acoes_verificacao.append(acao)

    def executar(self, auditado, debug=False):
        """Executa todas as ações, avalia a lógica do achado e retorna o achado, caso encontrado."""
        # Não é mais necessário fazer deepcopy aqui. A cópia será feita no nível do Auditado.

        [acao.executar(auditado, debug) for acao in self.acoes_verificacao]
        resultados = {acao.id: acao.resultado for acao in self.acoes_verificacao}

        # Avalia a lógica do achado com os resultados das ações
        achado_ocorreu = avalia_logica(self.logica_achado, resultados)
        self.achado_ocorreu = achado_ocorreu

        self.executado = True

        if debug:
            # [print(acao) for acao in acoes_verificadas.values()]
            [print(f'{av}: {resultado}') for av, resultado in resultados.items()]
            print(f"Procedimento resultou em achado: {'Sim' if achado_ocorreu else 'Não'}")
            print()

        if achado_ocorreu:
            achado = Achado(numero=self.numero_achado, nome=self.nome_achado)

            # Deduplicação preservando a ordem de inserção
            encaminhamentos = []
            seen_encaminhamentos = set()
            
            evidencias = []
            seen_evidencias = set()
            
            situacoes_encontradas = []
            seen_situacoes = set()

            for acao in self.acoes_verificacao:
                if acao.resultado:
                    # Encaminhamentos
                    if acao.encaminhamento and acao.tipo_encaminhamento:
                        enc_tuple = (acao.encaminhamento, acao.tipo_encaminhamento)
                        if enc_tuple not in seen_encaminhamentos:
                            seen_encaminhamentos.add(enc_tuple)
                            encaminhamentos.append({'encaminhamento': acao.encaminhamento, 'tipo': acao.tipo_encaminhamento})

                    # Evidências
                    if acao.descricao_evidencia and acao.descricao_evidencia not in seen_evidencias:
                        seen_evidencias.add(acao.descricao_evidencia)
                        evidencias.append(acao.descricao_evidencia)

                    # Situações Encontradas
                    if not pd.isna(acao.descricao_situacao_inconforme) and acao.descricao_situacao_inconforme not in seen_situacoes:
                        seen_situacoes.add(acao.descricao_situacao_inconforme)
                        situacoes_encontradas.append(acao.descricao_situacao_inconforme)

            achado.encaminhamentos = encaminhamentos
            achado.evidencias = evidencias
            achado.situacoes_encontradas = situacoes_encontradas
            self.achado = achado

        # Otimização: Remove a referência ao DataFrame de todas as fontes de informação
        # usadas neste procedimento APÓS a execução de todas as ações.
        for acao in self.acoes_verificacao:
            if hasattr(acao.fonte_informacao, 'info') and acao.fonte_informacao.info is not None:
                acao.fonte_informacao.info = None

        return self

    def to_dict(self):
        return {
            'id': self.id,
            'descricao': self.descricao,
            'logica_achado': self.logica_achado,
            'numero_achado': self.numero_achado,
            'nome_achado': self.nome_achado,
            'executado': safe_serialize(self.executado),
            'acoes_verificacao': [acao.to_dict() for acao in self.acoes_verificacao],
            'achado': self.achado.to_dict() if self.achado else None,
            'achado_ocorreu': safe_serialize(self.achado_ocorreu)
        }

    @classmethod
    def from_dict(cls, data):
        obj = cls(
            id=data.get('id'),
            descricao=data.get('descricao'),
            logica_achado=data.get('logica_achado'),
            numero_achado=data.get('numero_achado'),
            nome_achado=data.get('nome_achado')
        )
        obj.executado = data.get('executado')
        obj.acoes_verificacao = [AcaoVerificacao.from_dict(av) for av in data.get('acoes_verificacao', [])]
        obj.achado = Achado.from_dict(data.get('achado')) if data.get('achado') else None
        obj.achado_ocorreu = data.get('achado_ocorreu')
        return obj

class Auditado:
    contador = 1  # Contador de instâncias para automatizar o identificador

    def __init__(self, nome, sigla, id=None):
        if id is None:
            id = f"A{Auditado.contador:02d}"
            Auditado.contador += 1  # Incrementa o contador para o próximo identificador

        self.id = id
        self.nome = nome
        self.sigla = sigla
        self.foi_auditado = False

        # Armazena os resultados dos procedimentos de auditoria
        self.procedimentos_executados = []
        self.tem_achados = False

    def __repr__(self):
        return (f"Auditado(id='{self.id}', sigla='{self.sigla}')\n" +
                f"nome='{self.nome}'\n" +
                f"foi_auditado='{self.foi_auditado}'\n" +
                f"tem_achados='{self.tem_achados}'\n")

    def __aplicar_procedimento(self, procedimento, debug=False):
        if debug:
            print(f'Aplicando procedimento {procedimento.id} em {self.sigla}')
            print(f'Em busca do achado {procedimento.nome_achado}')
            print(f'Lógica {procedimento.logica_achado}')
            print()

        if procedimento.id in [p.id for p in self.procedimentos_executados]:
            # print(f'Procedimento {procedimento.id} já foi executado')
            return

        # Cria uma cópia do procedimento AQUI, uma vez por auditado.
        p = copy.deepcopy(procedimento)
        p.executar(self.sigla, debug)
        self.procedimentos_executados.append(p)

        if p.achado:
            self.tem_achados = True

    def aplicar_procedimentos(self, procedimentos, debug=False):
        for procedimento in procedimentos: # procedimentos é uma lista de objetos originais
            self.__aplicar_procedimento(procedimento, debug)

        self.foi_auditado = True

    def show(self):
        """Retorna uma string formatada com os dados do auditado."""
        report_lines = []
        report_lines.append("==================================================")
        report_lines.append(f"            Relatório do Auditado - {self.sigla}            ")
        report_lines.append("==================================================")
        report_lines.append(f"Sigla: {self.sigla}")
        report_lines.append(f"Nome: {self.nome}")
        report_lines.append(f"Foi auditado: {'Sim' if self.foi_auditado else 'Não'}")
        report_lines.append(f"Tem achados: {'Sim' if self.tem_achados else 'Não'}")
        report_lines.append("\n--- Procedimentos Aplicados ---")
        if self.procedimentos_executados:
            for p in self.procedimentos_executados:
                achado_info = f"(Achado: {p.achado.nome})" if p.achado else "(Sem achado)"
                report_lines.append(f"  - {p.id}: {p.descricao} {achado_info}")
        else:
            report_lines.append("  Nenhum procedimento aplicado ainda.")
        report_lines.append("\n--- Lista de Achados Encontrados ---")
        nomes_achados = self.get_nomes_achados()
        if nomes_achados:
            for achado in nomes_achados:
                report_lines.append(f"  - {achado}")
        else:
            report_lines.append("  Nenhum achado encontrado.")
        report_lines.append("==================================================\n")

        return "\n".join(report_lines)

    def get_nomes_achados(self):
        """Retorna uma lista dos nomes dos achados identificados para o auditado."""
        return [f"{p.achado.numero}. {p.achado.nome}" for p in self.procedimentos_executados if p.achado is not None]

    def get_achados(self):
        """Retorna uma lista dos nomes dos achados identificados para o auditado."""
        return {f"achado{p.achado.numero}": p.achado for p in self.procedimentos_executados if p.achado is not None}

    def get_achado_por_nome(self, nome_achado):
        """Retorna o objeto Achado correspondente ao nome fornecido."""
        for p in self.procedimentos_executados:
            if p.achado and p.achado.nome == nome_achado:
                return p.achado
        return None

    def get_situacoes_inconformes(self):
        situacoes = []
        for p in self.procedimentos_executados:
            if p.achado is not None:
                for s in p.achado.situacoes_encontradas:
                    situacoes.append(s.strip())
        return situacoes

    def get_encaminhamentos(self):
        """Retorna uma lista de todos os encaminhamentos aplicados ao auditado."""
        encaminhamentos = []
        for p in self.procedimentos_executados:
            if p.achado is not None:
                for e in p.achado.encaminhamentos:
                    encaminhamentos.append(e['encaminhamento'].strip())

        return list(set(encaminhamentos))

    def get_plano_acao(self):
        """Retorna uma lista de todos os achados e os encaminhamentos sugeridos ao auditado."""

        encaminhamentos = []
        for p in self.procedimentos_executados:
            for acao in p.acoes_verificacao:
                if acao.resultado:
                    if len(encaminhamentos):
                        encontrou = False
                        for e in encaminhamentos:
                            if e['encaminhamento'] == acao.encaminhamento and e['tipo'] == acao.tipo_encaminhamento and e['achado_num'] == p.achado.numero:
                                encontrou = True

                        if not encontrou:
                            encaminhamentos.append({'achado_num': p.achado.numero, 'encaminhamento': acao.encaminhamento, 'tipo': acao.tipo_encaminhamento})
                    else:
                        encaminhamentos.append({'achado_num': p.achado.numero, 'encaminhamento': acao.encaminhamento, 'tipo': acao.tipo_encaminhamento})

        return encaminhamentos

    def reporta_procedimentos(self):
        conteudo_md = f"# {self.sigla} - {self.nome}\n\n"

        if self.tem_achados:
            conteudo_md += f"## Achados encontrados na organização\n"
            for p in self.procedimentos_executados:
                if p.achado:
                    conteudo_md += f"1. {p.achado.nome}\n"
                    if len(p.achado.situacoes_encontradas):
                        # conteudo_md += f"   **Situações encontradas**\n"
                        for s in p.achado.situacoes_encontradas:
                            conteudo_md += f"   - {s}\n"
            conteudo_md += f"\n"

        # Lista de procedimentos e achados
        conteudo_md += "## Procedimentos de Auditoria Aplicados\n"

        for idx, p in enumerate(self.procedimentos_executados):
            conteudo_md += f"### {idx+1}. Procedimento {p.id}\n"
            conteudo_md += f"- Descrição: {p.descricao}\n"
            conteudo_md += f"- Condição: {p.logica_achado}\n"
            conteudo_md += f"- Achado Materializado: {'Sim' if p.achado else 'Não'}\n"

            if p.achado:
                conteudo_md += f" - Nome do Achado: {p.achado.nome}\n"
                conteudo_md += f"  - **Evidências Encontradas**\n"

                # Evidências para cada ação de verificação que materializou o achado
                for a in p.acoes_verificacao:
                    if a.resultado:
                        conteudo_md += f"   - {a.descricao_evidencia}\n"

                conteudo_md += f"  - **Encaminhamentos propostos**\n"

                # Encaminhamentos propostos
                for a in p.acoes_verificacao:
                    if a.resultado:
                        conteudo_md += f"   - [{a.tipo_encaminhamento}] {a.encaminhamento}\n"

            # Detalhamento das ações de verificação para cada procedimento
            conteudo_md += f"#### {idx+1}.1. Ações de Verificação Aplicadas\n"
            for a in p.acoes_verificacao:
                conteudo_md += f" - **Ação {a.id}**\n"
                conteudo_md += f"  - Fonte de Informação: {a.fonte_informacao.descricao}\n"
                conteudo_md += f"  - Campo de Dados Buscado: {a.informacao_requerida}\n"
                conteudo_md += f"  - Situação Encontrada: {a.situacao_encontrada or 'Não encontrada'}\n"
                conteudo_md += f"  - Situação considerada como inconforme: {a.situacao_inconforme or 'Não encontrada'}\n"
                conteudo_md += f"  - Achado na Verificação: {'Sim' if a.resultado else 'Não'}\n"
                conteudo_md += "\n"

        return conteudo_md

    def documenta_procedimentos(self, template_path='docs/template_report.docx'):
        """
            Cria um documento .docx em memória com os dados do objeto Auditado, usando um template.
        """
        # Carregar o documento template
        doc = Document(template_path)

        # Pegar a coleção de estilos do documento
        styles = doc.styles

        # Aplicar justificado nos estilos que você usa
        # Adicione todos os estilos que você quer justificar
        try:
            styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            styles['List Bullet'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            styles['List Bullet 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            styles['List Bullet 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except KeyError as e:
            print(f"Aviso: Estilo {e} não encontrado no template. Ignorando justificação para ele.")

        # Título do relatório
        doc.add_heading(f"{self.sigla} - {self.nome}", level=1)

        if self.tem_achados:
            doc.add_heading("Achados encontrados na organização", level=2)
            for p in self.procedimentos_executados:
                if p.achado:
                    doc.add_paragraph(f"{p.achado.nome}", style="List Bullet")
            doc.add_paragraph(" ")

        # Lista de procedimentos de auditoria e achados
        doc.add_heading("Procedimentos de Auditoria Aplicados", level=2)

        for idx, p in enumerate(self.procedimentos_executados):
            # Detalhes do procedimento
            doc.add_heading(f"{idx+1}. Procedimento {p.id}", level=3)
            doc.add_paragraph(f"Descrição: {p.descricao}", style="List Bullet")
            doc.add_paragraph(f"Condição: {p.logica_achado}", style="List Bullet")
            doc.add_paragraph(f"Achado Materializado: {'Sim' if p.achado else 'Não'}", style="List Bullet")

            if p.achado:
                doc.add_paragraph(f"Nome do Achado: {p.achado.nome}", style="List Bullet")

                paragrafo_acao_id = doc.add_paragraph(style="List Bullet 2")
                run = paragrafo_acao_id.add_run("Evidências Encontradas")
                run.bold = True

                # Evidências para cada ação de verificação que materializou o achado
                for a in p.acoes_verificacao:
                    if a.resultado:
                        doc.add_paragraph(f"{a.descricao_evidencia}", style="List Bullet 3")

                paragrafo_acao_id = doc.add_paragraph(style="List Bullet 2")
                run = paragrafo_acao_id.add_run("Encaminhamentos propostos")
                run.bold = True

                # Encaminhamentos propostos
                for a in p.acoes_verificacao:
                    if a.resultado:
                        doc.add_paragraph(f"[{a.tipo_encaminhamento}] {a.encaminhamento}", style="List Bullet 3")

            # Ações de verificação detalhadas
            doc.add_heading(f"{idx+1}.1. Ações de Verificação Aplicadas", level=4)
            for a in p.acoes_verificacao:
                paragrafo_acao_id = doc.add_paragraph(style="Normal")
                run = paragrafo_acao_id.add_run(f"Ação {a.id}")
                run.bold = True

                doc.add_paragraph(f"Fonte de Informação: {a.fonte_informacao.descricao}", style="List Bullet")
                doc.add_paragraph(f"Campo de Dados Buscado: {a.informacao_requerida}", style="List Bullet")
                doc.add_paragraph(f"Situação Encontrada: {a.situacao_encontrada or 'Não encontrada'}", style="List Bullet")
                doc.add_paragraph(f"Situação considerada como inconforme: {a.situacao_inconforme or 'Não encontrada'}", style="List Bullet")
                doc.add_paragraph(f"Achado na Verificação: {'Sim' if a.resultado else 'Não'}", style="List Bullet")
                doc.add_paragraph(" ")
            #doc.add_paragraph(" ")  # Adiciona uma linha em branco entre procedimentos

        return doc

    def to_dict(self):
        return {
            'id': self.id,
            'nome': self.nome,
            'sigla': self.sigla,
            'foi_auditado': safe_serialize(self.foi_auditado),
            'tem_achados': safe_serialize(self.tem_achados),
            'procedimentos_executados': [p.to_dict() for p in self.procedimentos_executados]
        }

    @classmethod
    def from_dict(cls, data):
        obj = cls(
            id=data.get('id'),
            nome=data.get('nome'),
            sigla=data.get('sigla')
        )
        obj.foi_auditado = data.get('foi_auditado')
        obj.tem_achados = data.get('tem_achados')
        obj.procedimentos_executados = [ProcedimentoAuditoria.from_dict(p) for p in data.get('procedimentos_executados', [])]
        return obj

def gerar_tabela_achados(auditados):
    # Reconstrói o dicionário de procedimentos a partir dos achados em cada auditado
    procedimentos = {}
    # Pega os procedimentos aplicados em qualquer um:
    for p in list(auditados.values())[0].procedimentos_executados:
        procedimentos[p.id] = p

    # Coleta todos os nomes de achados únicos
    nomes_todos_achados = sorted({f"{p.numero_achado}. {p.nome_achado}" for p in procedimentos.values()})

    # Inicializa uma lista para armazenar os dados da tabela
    dados_tabela = []

    # verifica quais achados foram encontrados
    for auditado in auditados.values():
        if auditado.foi_auditado:
            achados_auditado = auditado.get_nomes_achados()
            linha = {achado: ('X' if achado in achados_auditado else '') for achado in nomes_todos_achados}
            linha["Auditado"] = auditado.sigla  # Adiciona o identificador do auditado
            dados_tabela.append(linha)
        else:
            print(f'{auditado.sigla} ainda não foi auditado')

    # Cria o DataFrame com os dados coletados
    df_achados = pd.DataFrame(dados_tabela)
    df_achados = df_achados.set_index("Auditado")

    return df_achados

def gerar_tabela_encaminhamentos(auditados):
    # Reconstrói o dicionário de procedimentos a partir dos achados em cada auditado
    procedimentos = {}
    # Pega os procedimentos aplicados em qualquer um:
    for p in list(auditados.values())[0].procedimentos_executados:
        procedimentos[p.id] = p


    # Coleta todos os encaminhamentos únicos
    todos_encaminhamentos = sorted({(acao.tipo_encaminhamento, acao.encaminhamento) for p in procedimentos.values()
                                    for acao in p.acoes_verificacao if acao.encaminhamento}, key=lambda x: (x[0], x[1]))

    # Inicializa uma lista para armazenar os dados da tabela
    dados_tabela = []

    # Aplica os procedimentos a cada auditado e verifica quais encaminhamentos foram encontrados
    for auditado in auditados.values():
        if auditado.foi_auditado:
            encaminhamentos_auditado = auditado.get_encaminhamentos()

            linha = {f"[{tipo}] {encaminhamento}": ('X' if encaminhamento in encaminhamentos_auditado else '') for (tipo, encaminhamento) in todos_encaminhamentos}

            linha["Auditado"] = auditado.sigla  # Adiciona o identificador do auditado
            dados_tabela.append(linha)
        else:
            print(f'{auditado.sigla} ainda não foi auditado')

    # Cria o DataFrame com os dados coletados
    df_encaminhamentos = pd.DataFrame(dados_tabela)
    df_encaminhamentos = df_encaminhamentos.set_index("Auditado")

    return df_encaminhamentos

def gerar_tabela_situacoes_inconformes(auditados):
    # Reconstrói o dicionário de procedimentos a partir dos achados em cada auditado
    procedimentos = {}
    # Pega os procedimentos aplicados em qualquer um:
    for p in list(auditados.values())[0].procedimentos_executados:
        procedimentos[p.id] = p


    # Coleta todos os encaminhamentos únicos
    todas_situacoes_inconformes = []
    seen_situacoes = set()
    for p in procedimentos.values():
        for acao in p.acoes_verificacao:
            texto = f"[ACHADO {p.numero_achado}] {acao.descricao_situacao_inconforme}"
            if texto not in seen_situacoes:
                seen_situacoes.add(texto)
                todas_situacoes_inconformes.append(texto)

    # todas_situacoes_inconformes = sorted({f"[ACHADO {p.numero_achado}] {acao.descricao_situacao_inconforme}" for p in procedimentos.values()
    #                                 for acao in p.acoes_verificacao})

    # Inicializa uma lista para armazenar os dados da tabela
    dados_tabela = []

    for auditado in auditados.values():
        if auditado.foi_auditado:
            situacoes_auditado = auditado.get_situacoes_inconformes()

            linha = {f"{situacao}": ('X' if any([s in situacao for s in situacoes_auditado]) else '') for situacao in todas_situacoes_inconformes}

            linha["Auditado"] = auditado.sigla  # Adiciona o identificador do auditado
            dados_tabela.append(linha)
        else:
            print(f'{auditado.sigla} ainda não foi auditado')

    # Cria o DataFrame com os dados coletados
    df_situacoes = pd.DataFrame(dados_tabela)
    df_situacoes = df_situacoes.set_index("Auditado")

    return df_situacoes